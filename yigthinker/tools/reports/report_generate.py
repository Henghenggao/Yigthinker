from __future__ import annotations
from pathlib import Path
from typing import Literal
import pandas as pd
from pydantic import BaseModel
from yigthinker.types import ToolResult
from yigthinker.session import SessionContext

ReportFormat = Literal["excel", "pdf", "csv", "docx"]


def _safe_output_path(
    output_path: str,
    ctx_settings: dict,
    attachments: set[Path] | None = None,
) -> tuple[Path, str | None]:
    """Return (resolved_path, error_msg). error_msg is None if safe.

    A path outside workspace_dir is accepted if it is present in the
    ``attachments`` allowlist (caller's contract: set entries are expected to
    be pre-resolved absolute paths). Symmetrical with df_load._safe_path so
    adapters that register a temp file can both read from and write to it.
    """
    workspace = Path(ctx_settings.get("workspace_dir", Path.cwd())).expanduser().resolve()
    raw_path = Path(output_path).expanduser()
    try:
        candidate = raw_path if raw_path.is_absolute() else workspace / raw_path
        resolved = candidate.resolve(strict=False)
    except Exception:
        return raw_path, f"Cannot resolve output path: {output_path}"
    try:
        resolved.relative_to(workspace)
        return resolved, None
    except ValueError:
        if attachments and resolved in attachments:
            return resolved, None
        return resolved, (
            f"Access denied: output path '{output_path}' is outside workspace '{workspace}'."
        )


class ReportGenerateInput(BaseModel):
    var_name: str = "last_query"
    format: ReportFormat = "excel"
    output_path: str
    title: str = "Report"
    sheet_name: str = "Data"


class ReportGenerateTool:
    name = "report_generate"
    description = (
        "Use this when the user asks for an Excel / PDF / CSV / DOCX file "
        "from a DataFrame that's already in the registry. Prefer this over "
        "explaining how to make a report — just generate it. "
        "Excel output uses openpyxl with formatted headers. "
        "PDF output uses reportlab for tabular layout. "
        "DOCX output uses python-docx with a styled Light Grid Accent 1 table. "
        "Returns the output path in the result; the channel adapter will "
        "deliver it as a file card."
    )
    input_schema = ReportGenerateInput

    async def execute(self, input: ReportGenerateInput, ctx: SessionContext) -> ToolResult:
        try:
            df = ctx.vars.get(input.var_name)
        except KeyError as exc:
            return ToolResult(tool_use_id="", content=str(exc), is_error=True)

        path, err = _safe_output_path(
            input.output_path, ctx.settings, attachments=ctx.attachments
        )
        if err:
            return ToolResult(tool_use_id="", content=err, is_error=True)
        path.parent.mkdir(parents=True, exist_ok=True)

        try:
            if input.format == "excel":
                self._write_excel(df, path, input.title, input.sheet_name)
            elif input.format == "csv":
                df.to_csv(path, index=False)
            elif input.format == "pdf":
                self._write_pdf(df, path, input.title)
            elif input.format == "docx":
                self._write_docx(df, path, input.title)

            return ToolResult(
                tool_use_id="",
                content={
                    "output_path": str(path),
                    "format": input.format,
                    "rows": len(df),
                    "columns": len(df.columns),
                },
            )
        except Exception as exc:
            return ToolResult(tool_use_id="", content=str(exc), is_error=True)

    def _write_excel(self, df: pd.DataFrame, path: Path, title: str, sheet_name: str) -> None:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
        from openpyxl.utils import get_column_letter

        wb = openpyxl.Workbook()
        ws = wb.active
        ws.title = sheet_name

        # Title row
        ws.merge_cells(f"A1:{get_column_letter(len(df.columns))}1")
        title_cell = ws["A1"]
        title_cell.value = title
        title_cell.font = Font(bold=True, size=14)
        title_cell.alignment = Alignment(horizontal="center")

        # Header row
        header_fill = PatternFill(start_color="1E40AF", end_color="1E40AF", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF")
        for col_idx, col_name in enumerate(df.columns, 1):
            cell = ws.cell(row=2, column=col_idx, value=col_name)
            cell.fill = header_fill
            cell.font = header_font
            ws.column_dimensions[get_column_letter(col_idx)].width = max(12, len(str(col_name)) + 4)

        # Data rows
        for row_idx, row in enumerate(df.itertuples(index=False), 3):
            for col_idx, value in enumerate(row, 1):
                ws.cell(row=row_idx, column=col_idx, value=value)

        wb.save(path)

    def _write_pdf(self, df: pd.DataFrame, path: Path, title: str) -> None:
        from reportlab.lib.pagesizes import A4, landscape
        from reportlab.platypus import SimpleDocTemplate, LongTable, TableStyle, Paragraph
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors

        doc = SimpleDocTemplate(str(path), pagesize=landscape(A4))
        styles = getSampleStyleSheet()
        elements = [Paragraph(title, styles["Title"])]

        data = [list(df.columns)] + df.values.tolist()
        # LongTable with repeatRows=1 paginates large tables correctly and
        # repeats the header row on each page, unlike plain Table which can
        # overflow the frame or omit the header on continuation pages.
        table = LongTable(data, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1E40AF")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#DBEAFE")]),
        ]))
        elements.append(table)
        doc.build(elements)

    def _write_docx(self, df: pd.DataFrame, path: Path, title: str) -> None:
        from docx import Document
        from docx.enum.table import WD_TABLE_ALIGNMENT

        doc = Document()
        doc.add_heading(title, level=1)

        table = doc.add_table(
            rows=len(df) + 1,
            cols=len(df.columns),
            style="Light Grid Accent 1",
        )
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header row
        for col_idx, col_name in enumerate(df.columns):
            cell = table.rows[0].cells[col_idx]
            cell.text = str(col_name)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        # Data rows
        for row_idx, row in enumerate(df.itertuples(index=False), 1):
            for col_idx, value in enumerate(row):
                table.rows[row_idx].cells[col_idx].text = str(value)

        doc.save(str(path))
