"""Tests for DOCX report format (Task 18)."""
from __future__ import annotations

import pandas as pd
import pytest

from yigthinker.session import SessionContext
from yigthinker.tools.reports.report_generate import (
    ReportGenerateInput,
    ReportGenerateTool,
)


@pytest.mark.asyncio
async def test_docx_report_generates_file(tmp_path):
    ctx = SessionContext(settings={"workspace_dir": str(tmp_path)})
    df = pd.DataFrame({"product": ["A", "B"], "revenue": [100, 200]})
    ctx.vars.set("data", df)

    tool = ReportGenerateTool()
    inp = ReportGenerateInput(
        var_name="data",
        format="docx",
        output_path=str(tmp_path / "report.docx"),
        title="Revenue Report",
    )
    result = await tool.execute(inp, ctx)
    assert not result.is_error, f"DOCX generation failed: {result.content}"
    out = tmp_path / "report.docx"
    assert out.exists()
    assert out.stat().st_size > 1000
    # Docx files are zip archives; verify magic bytes
    with out.open("rb") as f:
        assert f.read(2) == b"PK"


@pytest.mark.asyncio
async def test_docx_report_contains_title_and_rows(tmp_path):
    """Verify the rendered document actually contains the title heading and
    all data rows by re-reading the DOCX.
    """
    from docx import Document

    ctx = SessionContext(settings={"workspace_dir": str(tmp_path)})
    df = pd.DataFrame({
        "name": ["Alice", "Bob", "Carol"],
        "score": [10, 20, 30],
    })
    ctx.vars.set("data", df)

    tool = ReportGenerateTool()
    out_path = tmp_path / "scores.docx"
    inp = ReportGenerateInput(
        var_name="data",
        format="docx",
        output_path=str(out_path),
        title="Scoreboard",
    )
    result = await tool.execute(inp, ctx)
    assert not result.is_error

    doc = Document(str(out_path))
    headings = [p.text for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert "Scoreboard" in headings

    assert len(doc.tables) == 1
    table = doc.tables[0]
    # header row + 3 data rows
    assert len(table.rows) == 4
    # Header cells
    header = [c.text for c in table.rows[0].cells]
    assert header == ["name", "score"]
    # Data rows
    assert [c.text for c in table.rows[1].cells] == ["Alice", "10"]
    assert [c.text for c in table.rows[3].cells] == ["Carol", "30"]
